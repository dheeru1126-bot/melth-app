import os
import sys

# Install python-docx if not present
try:
    import docx
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_srs():
    doc = Document()
    
    # Title
    title = doc.add_heading('SOFTWARE REQUIREMENTS SPECIFICATION', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_heading('Mental Health Support Portal (Melth)', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Document Version: 2.0 (Updated with AI Chatbot & CBT Tools)')
    doc.add_paragraph('Standard: IEEE 830-1998')
    doc.add_paragraph('Date: March 2026')
    doc.add_paragraph('Target Platform: Web Application')
    doc.add_page_break()

    # Introduction
    doc.add_heading('1. Introduction', level=1)
    doc.add_heading('1.1 Purpose', level=2)
    doc.add_paragraph(
        "This Software Requirements Specification (SRS) document defines the complete functional and non-functional requirements for the Melth Mental Health Support Portal. "
        "Melth is a web-based, AI-driven application designed to support mental well-being through conversational AI, CBT-based exercises, and mood tracking. "
        "This document reflects the latest architecture shifts, including the transition to a Python/Flask backend, SQLite database, and integration with OpenAI."
    )
    
    doc.add_heading('1.2 Scope', level=2)
    doc.add_paragraph(
        "The portal is designed for the general public. It empowers users to monitor their mental health, practice clinically-informed Cognitive Behavioral Therapy (CBT) techniques, and engage with a non-judgmental AI companion inspired by Woebot."
    )
    doc.add_paragraph("The system WILL:")
    ul_items = [
        "Provide an empathetic, Woebot-inspired AI Chatbot powered by OpenAI with CBT-informed conversational logic.",
        "Include a dedicated CBT Exercises Hub featuring Box Breathing, Thought Records, 5-4-3-2-1 Grounding, and Gratitude Journaling.",
        "Allow users to log their mood daily and visualize trends via a personalized graphical dashboard.",
        "Detect crisis keywords (self-harm, suicide) automatically during AI chats to intervene with emergency resources like the 988 Lifeline.",
        "Ensure data privacy and present a strong safety disclaimer before users access AI support."
    ]
    for item in ul_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph("\nThe system will NOT:")
    ul_items_not = [
        "Provide clinical diagnoses or replace licensed human therapists.",
        "Handle psychiatric emergencies directly (beyond providing external helpline numbers).",
        "Sell or expose private emotional data to third-party data brokers."
    ]
    for item in ul_items_not:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('1.3 Key Technical Definitions', level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Term'
    hdr_cells[1].text = 'Definition'
    
    definitions = [
        ('CBT', 'Cognitive Behavioral Therapy - psychological treatment focusing on changing unhelpful cognitive distortions.'),
        ('OpenAI API', 'Cloud-based AI engine used to power the advanced conversational generation of the Melth Assistant.'),
        ('Flask', 'A lightweight Python web framework used to build the backend REST API.'),
        ('SQLite', 'A C-language library that implements a small, fast, self-contained SQL database engine used for local secure storage.'),
        ('React.js', 'Frontend JavaScript library used to create the interactive, single-page application (SPA).'),
        ('JWT', 'JSON Web Token — used for secure, stateless authentication between React and Flask.')
    ]
    for term, definition in definitions:
        row_cells = table.add_row().cells
        row_cells[0].text = term
        row_cells[1].text = definition

    doc.add_page_break()

    # Overall Description
    doc.add_heading('2. Overall Description', level=1)
    doc.add_heading('2.1 Product Features', level=2)
    
    features = [
        "F-01: AI CBT Assistant - A conversational companion utilizing advanced system prompts to provide empathetic, behavioral support and grounding.",
        "F-02: Rapid Crisis Detection - A safety mechanism that scans chat inputs for high-risk flags and instantly redirects to 988 emergency protocols.",
        "F-03: Dashboard & Mood Analytics - Charts displaying emotional fluctuations alongside actionable 'Melth AI Insights' derived from recent entries.",
        "F-04: Structured CBT Exercises - Interactive web modules teaching Box Breathing, Thought Reconstruction, and Grounding.",
        "F-05: User Authentication - Secure registration, login, and session preservation using JWT and Flask-Bcrypt.",
        "F-06: Resource Library - Categorized self-help literature and breathing/sleep techniques."
    ]
    for f in features:
        doc.add_paragraph(f, style='List Bullet')

    doc.add_heading('2.2 Operating Environment & Technologies', level=2)
    env_items = [
        "Frontend: React.js, Tailwind CSS, Recharts, Lucide-React.",
        "Backend: Python 3, Flask, Flask-SQLAlchemy, Flask-JWT-Extended.",
        "AI Engine: OpenAI gpt-3.5/4 models via REST, with a local rule-based fallback mechanism.",
        "Database: local SQLite (database.db) via SQLAlchemy ORM.",
        "Browser: Chrome, Safari, Firefox, Edge."
    ]
    for e in env_items:
        doc.add_paragraph(e, style='List Bullet')

    doc.add_heading('2.3 System Architecture Overview', level=2)
    doc.add_paragraph(
        "The system relies on a modern decoupled architecture. The React SPA communicates with the Python Flask endpoint via Axios over HTTP. "
        "The Python backend manages the SQLite database directly through SQLAlchemy models (User, Mood, Chat). "
        "When a user talks to the AI, the backend intercepts the message, adds a strict CBT system prompt context, sends it to the OpenAI API, and returns the response to the user interface in a smooth, typing-animated chat bubble."
    )

    doc.add_heading('3. Key Interface Upgrades', level=1)
    doc.add_paragraph(
        "Aesthetic constraints demand a premium, calming design language. The UI was heavily updated to include:"
    )
    ui_items = [
        "Soft pastel color palettes (Baby blues, greens, soft indigo defaults to #f8fafc).",
        "Large rounded corner radii (rounded-2xl to rounded-[2rem]) for a friendlier application feel.",
        "Micro-animations including fade-ins and typing indicators inside the chat console.",
        "Mandatory transparent clinical safety disclaimers prior to accessing AI support."
    ]
    for ui in ui_items:
        doc.add_paragraph(ui, style='List Bullet')

    doc.add_page_break()
    
    doc.add_heading('Executive Summary Update', level=1)
    summary = doc.add_paragraph(
        "This version of the SRS marks the pivotal shift of the Mental Health Support Portal towards specialized AI companionship. "
        "By shedding the generic tracker model and implementing robust, clinical-style CBT interventions and real-time OpenAI conversational support, "
        "the application operates on a profoundly more engaging and useful tier. The backend has been vastly simplified into a Python/Flask structure to better "
        "suit Native AI integrations, making deployment much faster while heavily improving user data privacy by confining logs directly within SQLite isolated environments."
    )

    doc.save('Melth_SRS_Document_V2.docx')

if __name__ == '__main__':
    create_srs()
